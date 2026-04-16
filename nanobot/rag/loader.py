"""Legal document loader — PDF, HTML, DOCX, and plain text."""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from loguru import logger


@dataclass
class RawDocument:
    """A parsed legal document with its content and metadata."""

    title: str
    text: str
    source_path: str  # file path or URL
    doc_type: str = "law"  # law / judicial_interpretation / case / contract_template
    law_area: str = ""
    effective_date: str = ""
    extra: dict = field(default_factory=dict)


class LegalDocumentLoader:
    """法律文档加载器 — 解析 PDF/HTML/DOCX/TXT 格式的法律文档."""

    # Patterns to extract metadata from legal text
    _TITLE_PATTERN = re.compile(r"《([^》]+)》")
    _LAW_AREA_KEYWORDS: dict[str, list[str]] = {
        "民法": ["民法典", "民法", "合同", "物权", "侵权", "婚姻", "继承", "人格权"],
        "刑法": ["刑法", "刑事", "犯罪", "刑罚"],
        "商法": ["公司法", "证券法", "破产法", "保险法", "票据法", "商法"],
        "劳动法": ["劳动法", "劳动合同", "劳动", "社保", "工伤保险"],
        "行政法": ["行政", "行政处罚", "行政许可", "行政复议"],
        "诉讼法": ["诉讼", "民事诉讼法", "刑事诉讼法", "行政诉讼法", "仲裁"],
        "知识产权": ["专利", "商标", "著作权", "知识产权"],
        "宪法": ["宪法"],
    }

    def load_file(self, path: Path) -> list[RawDocument]:
        """Load a legal document from a file, dispatching by extension."""
        path = Path(path)
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self.load_pdf(path)
        if suffix in (".html", ".htm"):
            return self.load_html_file(path)
        if suffix == ".docx":
            return self.load_docx(path)
        if suffix in (".txt", ".md"):
            return self.load_text(path)
        logger.warning("Unsupported file format: {}", suffix)
        return []

    def load_directory(self, data_dir: Path) -> list[RawDocument]:
        """Load all supported documents from a directory recursively."""
        docs: list[RawDocument] = []
        data_dir = Path(data_dir)
        if not data_dir.exists():
            logger.warning("Data directory does not exist: {}", data_dir)
            return []
        extensions = {".pdf", ".html", ".htm", ".docx", ".txt", ".md"}
        for fpath in sorted(data_dir.rglob("*")):
            if fpath.is_file() and fpath.suffix.lower() in extensions:
                try:
                    loaded = self.load_file(fpath)
                    docs.extend(loaded)
                except Exception as e:
                    logger.error("Failed to load {}: {}", fpath, e)
        logger.info("Loaded {} documents from {}", len(docs), data_dir)
        return docs

    def load_pdf(self, path: Path) -> list[RawDocument]:
        """Parse a legal PDF using PyMuPDF."""
        path = Path(path)
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.error("PyMuPDF not installed. Install with: pip install PyMuPDF")
            return []

        doc = fitz.open(str(path))
        pages_text = []
        for page in doc:
            pages_text.append(page.get_text())
        full_text = "\n".join(pages_text)
        doc.close()

        if not full_text.strip():
            return []

        metadata = self._infer_metadata(full_text, str(path))
        return [RawDocument(
            title=metadata.get("title", path.stem),
            text=self._clean_text(full_text),
            source_path=str(path),
            doc_type=metadata.get("doc_type", self._infer_doc_type(path)),
            law_area=metadata.get("law_area", ""),
            effective_date=metadata.get("effective_date", ""),
        )]

    def load_html_file(self, path: Path) -> list[RawDocument]:
        """Parse an HTML file, extracting text content."""
        path = Path(path)
        try:
            from readability import Document
            import httpx
        except ImportError:
            # Fallback to basic HTML stripping
            text = self._strip_html_tags(path.read_text(encoding="utf-8", errors="ignore"))
        else:
            html = path.read_text(encoding="utf-8", errors="ignore")
            try:
                doc = Document(html)
                text = self._strip_html_tags(doc.summary())
            except Exception:
                text = self._strip_html_tags(html)

        if not text.strip():
            return []

        metadata = self._infer_metadata(text, str(path))
        return [RawDocument(
            title=metadata.get("title", path.stem),
            text=self._clean_text(text),
            source_path=str(path),
            doc_type=metadata.get("doc_type", self._infer_doc_type(path)),
            law_area=metadata.get("law_area", ""),
            effective_date=metadata.get("effective_date", ""),
        )]

    async def load_html_url(self, url: str) -> list[RawDocument]:
        """Fetch and parse an HTML page from URL."""
        try:
            import httpx
        except ImportError:
            logger.error("httpx not installed")
            return []

        try:
            async with httpx.AsyncClient(timeout=30, follow_redirects=True) as client:
                resp = await client.get(url)
                resp.raise_for_status()
                html = resp.text
        except Exception as e:
            logger.error("Failed to fetch {}: {}", url, e)
            return []

        try:
            from readability import Document
            doc = Document(html)
            text = self._strip_html_tags(doc.summary())
            title = doc.title()
        except Exception:
            text = self._strip_html_tags(html)
            title = ""

        if not text.strip():
            return []

        metadata = self._infer_metadata(text, url)
        return [RawDocument(
            title=title or metadata.get("title", ""),
            text=self._clean_text(text),
            source_path=url,
            doc_type=metadata.get("doc_type", "law"),
            law_area=metadata.get("law_area", ""),
            effective_date=metadata.get("effective_date", ""),
        )]

    def load_docx(self, path: Path) -> list[RawDocument]:
        """Parse a DOCX file (contract templates, etc.)."""
        path = Path(path)
        try:
            from docx import Document as DocxDocument
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return []

        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)

        if not full_text.strip():
            return []

        return [RawDocument(
            title=path.stem,
            text=self._clean_text(full_text),
            source_path=str(path),
            doc_type="contract_template",
        )]

    def load_text(self, path: Path) -> list[RawDocument]:
        """Parse a plain text or markdown file."""
        path = Path(path)
        try:
            text = path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            text = path.read_text(encoding="gbk", errors="ignore")

        if not text.strip():
            return []

        metadata = self._infer_metadata(text, str(path))
        return [RawDocument(
            title=metadata.get("title", path.stem),
            text=self._clean_text(text),
            source_path=str(path),
            doc_type=metadata.get("doc_type", self._infer_doc_type(path)),
            law_area=metadata.get("law_area", ""),
            effective_date=metadata.get("effective_date", ""),
        )]

    def _infer_metadata(self, text: str, source: str) -> dict:
        """Infer metadata (title, law_area, effective_date) from text content."""
        meta: dict = {}

        # Try to extract law title from 《》
        title_match = self._TITLE_PATTERN.search(text[:2000])
        if title_match:
            meta["title"] = title_match.group(1)

        # Infer law_area from content keywords
        text_lower = text[:5000]
        for area, keywords in self._LAW_AREA_KEYWORDS.items():
            if any(kw in text_lower for kw in keywords):
                meta["law_area"] = area
                break

        # Infer doc_type from source path or content
        source_lower = source.lower()
        if "解释" in source_lower or "解释" in text[:2000]:
            meta["doc_type"] = "judicial_interpretation"
        elif "案例" in source_lower or "判决" in text[:2000] or "裁定" in text[:2000]:
            meta["doc_type"] = "case"
        elif "合同" in source_lower or "template" in source_lower:
            meta["doc_type"] = "contract_template"

        # Try to extract effective date
        date_pattern = re.compile(r"(\d{4})年(\d{1,2})月(\d{1,2})日")
        date_match = date_pattern.search(text[:3000])
        if date_match:
            y, m, d = date_match.groups()
            meta["effective_date"] = f"{y}-{int(m):02d}-{int(d):02d}"

        return meta

    def _infer_doc_type(self, path: Path) -> str:
        """Infer document type from file path."""
        parts = str(path).lower()
        if "司法解释" in parts or "interpretation" in parts:
            return "judicial_interpretation"
        if "案例" in parts or "case" in parts:
            return "case"
        if "合同" in parts or "contract" in parts or "template" in parts:
            return "contract_template"
        return "law"

    @staticmethod
    def _strip_html_tags(html: str) -> str:
        """Remove HTML tags and decode entities."""
        import html as html_mod
        text = re.sub(r"<[^>]+>", "", html)
        text = html_mod.unescape(text)
        # Normalize whitespace
        text = re.sub(r"\s+", " ", text)
        return text.strip()

    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean up extracted text: normalize whitespace, remove noise."""
        # Remove excessive blank lines
        text = re.sub(r"\n{3,}", "\n\n", text)
        # Remove common PDF artifacts
        text = re.sub(r"-\n", "", text)  # hyphenated line breaks
        return text.strip()
