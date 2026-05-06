"""Legal document loader — PDF, HTML, DOCX, and plain text."""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from loguru import logger


@dataclass
class RawDocument:
    """A parsed legal document with its content and metadata."""

    title: str
    text: str
    source_path: str  # file path or URL
    doc_type: str = "law"  # law / judicial_interpretation / case / contract_template
    law_area: str = ""
    effective_date: str = ""
    extra: dict = field(default_factory=dict)


class LegalDocumentLoader:
    """法律文档加载器 — 解析 PDF/HTML/DOCX/TXT 格式的法律文档."""

    # Patterns to extract metadata from legal text
    _TITLE_PATTERN = re.compile(r"《([^》]+)》")
    _LAW_AREA_KEYWORDS: dict[str, list[str]] = {
        # 民法 — 民法典及民事基本制度
        "民法": [
            "民法典", "物权", "侵权", "婚姻", "继承", "人格权",
            "买卖合同", "借款合同", "租赁合同", "担保", "合同编",
            "收养", "民事主体", "民事法律行为",
        ],
        # 刑法 — 犯罪与刑罚
        "刑法": [
            "刑法", "刑事", "犯罪", "刑罚", "监狱", "禁毒",
            "反家庭暴力", "反恐怖主义", "反有组织犯罪", "反洗钱",
            "反电信网络诈骗", "反间谍", "反食品浪费",
            "社区矫正", "预防未成年人犯罪",
        ],
        # 商法 — 公司、证券、金融、票据、海商
        "商法": [
            "公司法", "合伙企业法", "证券法", "证券投资基金法",
            "期货和衍生品法", "票据法", "保险法", "海商法",
            "破产法", "商业银行法", "银行业监督管理法",
            "信托法", "电子商务法", "外商投资法",
        ],
        # 金融法 — 财税、银行、保险、证券
        "金融法": [
            "预算法", "税收征收管理", "增值税", "消费税",
            "车船税", "车辆购置税", "耕地占用税", "资源税",
            "环境保护税", "契税", "城市维护建设税", "烟叶税",
            "审计法", "反洗钱", "商业银行", "银行监督管理",
            "票据法", "证券", "基金", "保险法",
        ],
        # 行政法 — 行政处罚、行政许可、行政强制、行政复议、行政诉讼
        "行政法": [
            "行政处罚", "行政许可", "行政强制", "行政复议", "行政诉讼",
            "城乡规划", "建筑法", "土地管理", "房地产管理",
            "消防救援衔", "消防法", "治安管理处罚",
            "海关法", "出入境", "护照", "居民身份证",
            "网络安全", "数据安全", "密码法",
        ],
        # 诉讼法 — 民事/刑事诉讼、仲裁、司法协助
        "诉讼法": [
            "民事诉讼法", "刑事诉讼法", "仲裁法", "海事诉讼",
            "法律援助", "引渡", "国际刑事司法协助",
            "涉外民事关系法律适用",
        ],
        # 知识产权法 — 专利、商标、著作权
        "知识产权": [
            "专利法", "商标法", "著作权法", "知识产权",
            "专利", "商标", "著作权",
        ],
        # 宪法
        "宪法": ["宪法"],
        # 环境法 — 环境保护与污染防治
        "环境法": [
            "环境保护法", "大气污染防治", "水污染防治",
            "土壤污染防治", "固体废物", "噪声污染防治",
            "放射性污染防治", "海洋环境保护", "环境影响评价",
            "清洁生产", "循环经济", "可再生能源",
            "青藏高原生态保护", "核安全", "生物安全",
        ],
        # 自然资源法 — 土地、水、矿产、森林、草原等
        "自然资源法": [
            "土地管理法", "水法", "矿产资源法", "煤炭法",
            "石油天然气管道保护", "森林法", "草原法",
            "湿地保护", "渔业法", "水土保持", "防沙治沙",
            "黑土地保护", "海域使用管理", "海岛保护",
        ],
        # 交通航运法 — 铁路、公路、航空、水运
        "交通航运法": [
            "铁路法", "道路交通安全", "民用航空", "国防交通",
            "海上交通安全", "港口法", "航道法", "船舶吨税",
            "海商法", "海事诉讼", "海上", "航空法",
        ],
        # 社会法 — 劳动、社会保障、民生
        "社会法": [
            "劳动法", "劳动合同", "社会保险法", "职业病防治",
            "就业促进", "工会法", "老年人权益保障",
            "未成年人保护", "妇女权益保障", "残疾人保障",
            "消费者权益保护", "家庭教育促进", "慈善法",
            "药品管理法", "疫苗管理", "基本医疗卫生",
            "健康促进", "精神卫生", "无偿献血",
            "国境卫生检疫", "特种设备安全",
        ],
        # 教科文卫法 — 教育、科技、文化、卫生
        "教科文卫法": [
            "教育法", "职业教育法", "高等教育法", "民办教育促进",
            "科学技术进步", "科学技术普及", "学位法", "教师法",
            "测绘法", "文物保护", "非物质文化遗产", "档案法",
            "广告法", "电影产业促进", "旅游法", "体育法",
            "红十字会法", "母婴保健",
        ],
        # 国防军事法
        "国防军事法": [
            "国防动员", "国防教育", "现役军官", "预备役人员",
            "海警法", "国防交通", "驻外外交人员",
        ],
        # 公共安全法 — 消防、安全生产、应急管理（非刑法的公共安全类）
        "公共安全法": [
            "特种设备安全", "矿山安全", "突发事件应对",
            "防震减灾", "防洪法", "安全生产",
            "枪支管理", "传染病", "突发公共卫生",
        ],
    }

    def load_file(self, path: Path) -> list[RawDocument]:
        """Load a legal document from a file, dispatching by extension."""
        path = Path(path)
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self.load_pdf(path)
        if suffix in (".html", ".htm"):
            return self.load_html_file(path)
        if suffix == ".docx":
            return self.load_docx(path)
        if suffix in (".txt", ".md"):
            return self.load_text(path)
        logger.warning("Unsupported file format: {}", suffix)
        return []

    def load_directory(self, data_dir: Path) -> list[RawDocument]:
        """Load all supported documents from a directory recursively."""
        docs: list[RawDocument] = []
        data_dir = Path(data_dir)
        if not data_dir.exists():
            logger.warning("Data directory does not exist: {}", data_dir)
            return []
        extensions = {".pdf", ".html", ".htm", ".docx", ".txt", ".md"}
        for fpath in sorted(data_dir.rglob("*")):
            if fpath.is_file() and fpath.suffix.lower() in extensions:
                try:
                    loaded = self.load_file(fpath)
                    docs.extend(loaded)
                except Exception as e:
                    logger.error("Failed to load {}: {}", fpath, e)
        logger.info("Loaded {} documents from {}", len(docs), data_dir)
        return docs

    def load_pdf(self, path: Path) -> list[RawDocument]:
        """Parse a legal PDF using PyMuPDF."""
        path = Path(path)
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.error("PyMuPDF not installed. Install with: pip install PyMuPDF")
            return []

        doc = fitz.open(str(path))
        pages_text = []
        for page in doc:
            pages_text.append(page.get_text())
        full_text = "\n".join(pages_text)
        doc.close()

        if not full_text.strip():
            return []

        metadata = self._infer_metadata(full_text, str(path))
        return [RawDocument(
            title=metadata.get("title", path.stem),
            text=self._clean_text(full_text),
            source_path=str(path),
            doc_type=metadata.get("doc_type", self._infer_doc_type(path)),
            law_area=metadata.get("law_area", ""),
            effective_date=metadata.get("effective_date", ""),
        )]

    def load_html_file(self, path: Path) -> list[RawDocument]:
        """Parse an HTML file, extracting text content."""
        path = Path(path)
        try:
            from readability import Document
            import httpx
        except ImportError:
            # Fallback to basic HTML stripping
            text = self._strip_html_tags(path.read_text(encoding="utf-8", errors="ignore"))
        else:
            html = path.read_text(encoding="utf-8", errors="ignore")
            try:
                doc = Document(html)
                text = self._strip_html_tags(doc.summary())
            except Exception:
                text = self._strip_html_tags(html)

        if not text.strip():
            return []

        metadata = self._infer_metadata(text, str(path))
        return [RawDocument(
            title=metadata.get("title", path.stem),
            text=self._clean_text(text),
            source_path=str(path),
            doc_type=metadata.get("doc_type", self._infer_doc_type(path)),
            law_area=metadata.get("law_area", ""),
            effective_date=metadata.get("effective_date", ""),
        )]

    async def load_html_url(self, url: str) -> list[RawDocument]:
        """Fetch and parse an HTML page from URL."""
        try:
            import httpx
        except ImportError:
            logger.error("httpx not installed")
            return []

        try:
            async with httpx.AsyncClient(timeout=30, follow_redirects=True) as client:
                resp = await client.get(url)
                resp.raise_for_status()
                html = resp.text
        except Exception as e:
            logger.error("Failed to fetch {}: {}", url, e)
            return []

        try:
            from readability import Document
            doc = Document(html)
            text = self._strip_html_tags(doc.summary())
            title = doc.title()
        except Exception:
            text = self._strip_html_tags(html)
            title = ""

        if not text.strip():
            return []

        metadata = self._infer_metadata(text, url)
        return [RawDocument(
            title=title or metadata.get("title", ""),
            text=self._clean_text(text),
            source_path=url,
            doc_type=metadata.get("doc_type", "law"),
            law_area=metadata.get("law_area", ""),
            effective_date=metadata.get("effective_date", ""),
        )]

    def load_docx(self, path: Path) -> list[RawDocument]:
        """Parse a DOCX file (contract templates, etc.)."""
        path = Path(path)
        try:
            from docx import Document as DocxDocument
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return []

        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)

        if not full_text.strip():
            return []

        return [RawDocument(
            title=path.stem,
            text=self._clean_text(full_text),
            source_path=str(path),
            doc_type="contract_template",
        )]

    def load_text(self, path: Path) -> list[RawDocument]:
        """Parse a plain text or markdown file (UTF-8 encoding)."""
        path = Path(path)
        # Always UTF-8; replace undecodable bytes with placeholder to detect issues
        text = path.read_text(encoding="utf-8", errors="replace")

        if not text.strip():
            return []

        metadata = self._infer_metadata(text, str(path))
        return [RawDocument(
            title=metadata.get("title", path.stem),
            text=self._clean_text(text),
            source_path=str(path),
            doc_type=metadata.get("doc_type", self._infer_doc_type(path)),
            law_area=metadata.get("law_area", ""),
            effective_date=metadata.get("effective_date", ""),
        )]

    def _infer_metadata(self, text: str, source: str) -> dict:
        """Infer metadata (title, law_area, effective_date) from text content."""
        meta: dict = {}

        # Try to extract law title from 《》
        title_match = self._TITLE_PATTERN.search(text[:2000])
        if title_match:
            meta["title"] = title_match.group(1)

        # Infer law_area from content keywords
        # Score each area by the longest matching keyword — more specific wins.
        text_lower = text[:5000]
        best_area = ""
        best_len = 0
        for area, keywords in self._LAW_AREA_KEYWORDS.items():
            for kw in keywords:
                if kw in text_lower and len(kw) > best_len:
                    best_area = area
                    best_len = len(kw)
        if best_area:
            meta["law_area"] = best_area

        # Infer doc_type from source path or content
        source_lower = source.lower()
        if "解释" in source_lower or "解释" in text[:2000]:
            meta["doc_type"] = "judicial_interpretation"
        elif "案例" in source_lower or "判决" in text[:2000] or "裁定" in text[:2000]:
            meta["doc_type"] = "case"
        elif "合同" in source_lower or "template" in source_lower:
            meta["doc_type"] = "contract_template"

        # Try to extract effective date
        date_pattern = re.compile(r"(\d{4})年(\d{1,2})月(\d{1,2})日")
        date_match = date_pattern.search(text[:3000])
        if date_match:
            y, m, d = date_match.groups()
            meta["effective_date"] = f"{y}-{int(m):02d}-{int(d):02d}"

        return meta

    def _infer_doc_type(self, path: Path) -> str:
        """Infer document type from file path."""
        parts = str(path).lower()
        if "司法解释" in parts or "interpretation" in parts:
            return "judicial_interpretation"
        if "案例" in parts or "case" in parts:
            return "case"
        if "合同" in parts or "contract" in parts or "template" in parts:
            return "contract_template"
        return "law"

    @staticmethod
    def _strip_html_tags(html: str) -> str:
        """Remove HTML tags and decode entities."""
        import html as html_mod
        text = re.sub(r"<[^>]+>", "", html)
        text = html_mod.unescape(text)
        # Normalize whitespace
        text = re.sub(r"\s+", " ", text)
        return text.strip()

    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean legal text: remove HTML/markdown noise, normalize whitespace.

        Handles common noise in .txt/.md legal documents:
        - HTML comment separators: <!-- INFO END -->  （保留分隔符后的正文）
        - Markdown filename headers: # filename.txt / ## 标题
        - Residual HTML tags: <div>, <p>, etc.
        - Excessive blank lines and hyphenated line breaks
        """
        # 1. Remove <!-- INFO END --> (keep content after it — it's the real text)
        text = re.sub(r"<!--\s*INFO\s*END\s*-->", "\n", text, flags=re.IGNORECASE)

        # 2. Remove Markdown filename/header lines (# 标题, ## 标题, etc.)
        lines = text.split("\n")
        lines = [
            l for l in lines
            if not re.match(r"^#{1,6}\s+\S", l.strip())  # # 标题, ## 章节, etc.
        ]
        text = "\n".join(lines)

        # 3. Remove residual HTML tags
        text = re.sub(r"<[^>]+>", "", text)

        # 4. Remove URLs
        text = re.sub(r"https?://\S+", "", text)

        # 5. Normalize excessive blank lines
        text = re.sub(r"\n{3,}", "\n\n", text)

        # 6. Remove hyphenated line breaks (PDF artifact: "法-\n条" → "法律条")
        text = re.sub(r"-\n", "", text)

        return text.strip()
