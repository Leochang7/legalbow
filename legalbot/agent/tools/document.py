"""legal_document_generate — tool for generating legal documents from case facts."""

from __future__ import annotations

import json
import os
from pathlib import Path
from typing import TYPE_CHECKING, Any

from legalbot.agent.tools.base import Tool, tool_parameters
from legalbot.agent.tools.schema import StringSchema, tool_parameters_schema

if TYPE_CHECKING:
    from legalbot.document.generator import LegalDocumentGenerator


def _text_to_docx(text: str, filepath: str | Path) -> str:
    """Convert plain text document to a .docx file using python-docx."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    for line in text.splitlines():
        line = line.rstrip()
        if not line:
            doc.add_paragraph()
            continue

        # Check if it's a heading (starts with Chinese or special markers)
        is_heading = (
            line.startswith("### ")
            or line.startswith("## ")
            or line.startswith("# ")
            or line.startswith("【")
            or line.startswith("一、")
            or line.startswith("二、")
            or line.startswith("三、")
            or line.startswith("四、")
            or line.startswith("五、")
            or line.startswith("六、")
            or "：" in line
            and len(line) < 40
        )

        para = doc.add_paragraph()
        run = para.add_run(line)
        run.font.size = Pt(12)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii", "宋体"
        )
        run._element.rPr.rFonts.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi", "宋体"
        )
        run._element.rPr.rFonts.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "宋体"
        )
        if is_heading:
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)

    os.makedirs(os.path.dirname(str(filepath)) or ".", exist_ok=True)
    doc.save(str(filepath))
    return str(filepath)


@tool_parameters(
    tool_parameters_schema(
        doc_type=StringSchema(
            "文书类型：complaint(起诉状)/defense(答辩状)/agent_opinion(代理词)/appeal(上诉状)/enforcement(执行申请书)",
            enum=["complaint", "defense", "agent_opinion", "appeal", "enforcement"],
        ),
        case_facts=StringSchema(
            "案件事实描述",
            min_length=10,
            max_length=10000,
        ),
        extra_variables=StringSchema(
            "额外变量（JSON格式），用于覆盖或补充从案件事实中提取的信息",
            nullable=True,
        ),
        law_areas=StringSchema(
            "法律领域过滤（可选）：民法/刑法/商法/劳动法/行政法等",
            nullable=True,
        ),
        save_path=StringSchema(
            "文件保存路径（包含文件名，如 /path/to/起诉状.docx），留空则自动生成",
            nullable=True,
        ),
        required=["doc_type", "case_facts"],
    )
)
class LegalDocumentGenerateTool(Tool):
    """法律文书起草工具 — 根据案件事实生成法律文书（起诉状、答辩状等）。"""

    def __init__(self, generator: LegalDocumentGenerator):
        self._generator = generator

    @property
    def name(self) -> str:
        return "legal_document_generate"

    @property
    def description(self) -> str:
        return (
            "根据案件事实生成法律文书（.docx格式）。支持生成起诉状、答辩状、代理词、上诉状和执行申请书。"
            "系统会自动检索相关法律法规条文作为依据，并按中国法律文书规范格式输出。"
        )

    @property
    def read_only(self) -> bool:
        return True

    @property
    def exclusive(self) -> bool:
        return True

    async def execute(
        self,
        doc_type: str,
        case_facts: str,
        extra_variables: str | None = None,
        law_areas: str | None = None,
        save_path: str | None = None,
        **kwargs: Any,
    ) -> str:
        extra_vars: dict[str, Any] | None = None
        if extra_variables:
            try:
                extra_vars = json.loads(extra_variables)
            except json.JSONDecodeError:
                pass
        areas = [law_areas] if law_areas else None
        result = await self._generator.generate(
            doc_type=doc_type,
            case_facts=case_facts,
            extra_variables=extra_vars or {},
            law_areas=areas,
        )

        # If generator returned an error message (not a valid document), return as-is
        if result.startswith("不支持的文书类型") or result.startswith("文书生成失败"):
            return result

        # Determine save path
        doc_type_names = {
            "complaint": "起诉状",
            "defense": "答辩状",
            "agent_opinion": "代理词",
            "appeal": "上诉状",
            "enforcement": "执行申请书",
        }
        doc_name = doc_type_names.get(doc_type, doc_type)
        if save_path:
            filepath = Path(save_path)
        else:
            workspace = Path.home() / ".legalbot" / "documents"
            filepath = workspace / f"{doc_name}.docx"

        try:
            saved = _text_to_docx(result, filepath)
            response = f"法律文书已生成并保存至：\n{saved}\n\n如需使用，请打开以上文件路径。"
            return response
        except Exception as e:
            # Fallback: return text if docx generation fails
            return result + (
                f"\n\n[警告：.docx 文件生成失败：{e}]"
            )
